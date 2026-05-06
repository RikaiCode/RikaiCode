import os
import io
import json
import re
import markdown 
from datetime import datetime
from PIL import Image
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

from config import LOGO_PATH, SKIP_EXTENSIONS
from analysis import build_full_tree

def create_txt_content(files_dict):
    output = io.StringIO()
    output.write("RIKAI CODE EXPORT\n")
    output.write("==========================================\n\n")
    output.write(f"Note: The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}\n\n")
    output.write(build_full_tree(files_dict) + "\n\n")
    output.write("FILE CONTENTS\n")
    output.write("==========================================\n\n")
    for filename, content in files_dict.items():
        output.write(f"\n\n----- START: {filename} -----\n\n")
        output.write(content)
        output.write(f"\n\n----- END: {filename} -----\n")
    return output.getvalue()

def export_html(files_dict):
    html = "<html><head><style>body{font-family:monospace; background:#171717; color:#ededed;} .file{margin-bottom:20px; border-bottom:1px solid #333;} h3{color:#d3a0eb;}</style></head><body>"
    html += f"<h1>RikaiCode Export</h1>"
    html += f"<p><strong>Note:</strong> The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}</p>"
    html += f"<pre>{build_full_tree(files_dict)}</pre><hr>"
    for k, v in files_dict.items():
        html += f"<div class='file'><h3>{k}</h3><pre>{v}</pre></div>"
    html += "</body></html>"
    return html


def export_docx(files_dict):
    doc = Document()

    if os.path.exists(LOGO_PATH):
        try:
            doc.add_picture(LOGO_PATH, width=Inches(1.0))
        except:
            pass
    doc.add_heading('RikaiCode Repository Export', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Note: The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}")
    doc.add_paragraph(build_full_tree(files_dict))
    for filename, content in files_dict.items():
        doc.add_heading(filename, level=1)
        try: doc.add_paragraph(content)
        except: doc.add_paragraph("[Content error]")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_pdf(files_dict):
    class PDF(FPDF):
        def __init__(self):
            super().__init__()
            self.logo_path = LOGO_PATH
        
        def footer(self):
       
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            

            gen_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.cell(0, 10, f'Generated on {gen_date}', 0, 0, 'R')
            
            
            if os.path.exists(self.logo_path):

                self.image(self.logo_path, x=10, y=self.h - 12, h=5)

                self.set_xy(17, self.h - 12)
                self.cell(0, 5, 'RikaiCode', 0, 0, 'L')
            else:
                self.set_x(10)
                self.cell(0, 10, 'RikaiCode', 0, 0, 'L')

    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="RikaiCode Repository Export", ln=1, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 5, f"Note: The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}")
    pdf.ln(5)
    
    for filename, content in files_dict.items():
        pdf.set_font("Arial", 'B', size=14)
        pdf.cell(200, 10, txt=filename, ln=1)
        pdf.set_font("Arial", size=10)
        try:
    
            if len(content) > 5000:
                safe_content = content[:5000] + "\n... [TRUNCATED FOR PDF PREVIEW]"
            else:
                safe_content = content
                
            safe_content = safe_content.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, safe_content)
            pdf.ln(5)
        except: pass
    

    pdf_bytes = pdf.output(dest='S')
    

    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode('latin-1')
        
    file_stream = io.BytesIO(pdf_bytes)
    file_stream.seek(0)
    return file_stream

def export_json(files_dict):
    data = {
        "metadata": { 
            "generated_at": datetime.now().isoformat(), 
            "file_count": len(files_dict), 
            "architecture": build_full_tree(files_dict),
            "excluded_types": list(SKIP_EXTENSIONS)
        },
        "files": [{"name": k, "content": v} for k, v in files_dict.items()]
    }
    return json.dumps(data, indent=2)

def strip_markdown(text):
    """Removes markdown syntax to produce clean text for DOCX/PDF."""

    text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)

    text = re.sub(r'\*\*|\*|__|_', '', text)

    text = re.sub(r'`{1,3}', '', text)
    return text

def export_ai_response_docx(title, content):
    doc = Document()

    if os.path.exists(LOGO_PATH):
        try:
            doc.add_picture(LOGO_PATH, width=Inches(1.0))
        except:
            pass
    doc.add_heading('RikaiCode AI Analysis', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_heading(title, level=1)
    

    clean_content = strip_markdown(content)
    doc.add_paragraph(clean_content)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_ai_response_pdf(title, content):
    class PDF(FPDF):
        def __init__(self):
            super().__init__()
            self.logo_path = LOGO_PATH
        
        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            gen_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.cell(0, 10, f'Generated on {gen_date}', 0, 0, 'R')
            
            if os.path.exists(self.logo_path):
                self.image(self.logo_path, x=10, y=self.h - 12, h=5)
                self.set_xy(17, self.h - 12)
                self.cell(0, 5, 'RikaiCode', 0, 0, 'L')
            else:
                self.set_x(10)
                self.cell(0, 10, 'RikaiCode', 0, 0, 'L')

    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=16)
    pdf.cell(0, 10, txt=title, ln=1, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", size=12)
    

    clean_content = strip_markdown(content)
    safe_content = clean_content.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 5, safe_content)
    
    pdf_bytes = pdf.output(dest='S')
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode('latin-1')
    file_stream = io.BytesIO(pdf_bytes)
    file_stream.seek(0)
    return file_stream