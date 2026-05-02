import streamlit as st
import os
import base64
import io
import zipfile
import requests
import plotly.express as px
import pandas as pd
import time
import re
import json
import math
import gc  
from PIL import Image
from docx import Document
from fpdf import FPDF
from datetime import datetime, timedelta
from collections import Counter

st.set_page_config(
    page_title="RikaiCode",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed",
)


SKIP_EXTENSIONS = {'.pdf', '.txt', '.rtf', '.json', '.xlsx', '.csv'}
LOGO_PATH = "assets/logo.png" if os.path.exists("assets/logo.png") else "logo.png"

MAX_LINES_INTERACTIVE_PREVIEW = 50000


st.markdown("""
<style>

    html, body, [class*="css"] {

        font-family: 'Inter', sans-serif;
    }

    h1, h2, h3 { color: #ededed !important; font-weight: 700 !important; }
    .main .block-container { padding-top: 2rem; padding-bottom: 5rem; }
    .accent-text { color: #d3a0eb !important; }

   
    .stFileUploader {
        border: 1px dashed #d3a0eb !important;
        background-color: #1e1e1e !important;
        border-radius: 10px; padding: 20px;
    }
    

    .stButton > button {
        background-color: #d3a0eb !important;
        color: #171717 !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 10px 24px !important;
        font-weight: bold !important;
 
        width: 100%;
    }
    .stButton > button:hover {
        background-color: #b589c8 !important;
        transform: translateY(-2px);
       
    }

    
    .streamlit-expanderHeader {
        background-color: #1e1e1e !important;
        color: #ededed !important;
        border: 1px solid #333 !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
    }
    .streamlit-expanderContent {
        background-color: #1e1e1e !important;
        border: 1px solid #333 !important;
        border-top: none !important;
        border-radius: 0 0 8px 8px !important;
    }


    [data-testid="stMetric"] {
        background-color: #1e1e1e !important; border: 1px solid #333 !important;
        border-radius: 10px; padding: 15px !important;
    }
    [data-testid="stMetricLabel"] { 
        color: #a0a0a0 !important; 
        font-size: 1.1rem !important; 
    }
    [data-testid="stMetricValue"] { color: #d3a0eb !important; font-size: 1.8rem !important; }


    .grade-box {
        background-color: #1e1e1e;
        border: 2px solid #7c5e8a;
        border-radius: 15px;
        padding: 20px;
        text-align: center;
        margin-bottom: 20px;

    }
    .grade-letter {
        font-size: 4rem;
        font-weight: bold;
        color: #d3a0eb;
        line-height: 1;
    }
    .grade-reason {
        text-align: left;
        font-size: 0.9rem;
        color: #a0a0a0;
        margin-top: 10px;
        border-top: 1px solid #333;
        padding-top: 10px;
    }


    .arch-box {
       
        border: 0.5px solid #333;
       
        padding: 20px;
        border-radius: 8px;
        font-family: 'Courier New', monospace;
        white-space: pre;
        overflow-x: auto;
        color: #ededed;
    }


    .footer-custom {
        position: fixed; bottom: 0; left: 0; width: 100%;
        background-color: #1e1e1e; border-top: 1px solid #333;
        padding: 15px 20px; display: flex; justify-content: space-between;
        align-items: center; z-index: 9999; font-size: 14px;
    }
    .footer-link {
        color: #d3a0eb !important;
        text-decoration: none !important;
    }

    .footer-link:hover {
        color: #9d76ad !important; 

    }


    @keyframes float {
        0% { transform: translateY(0px); } 50% { transform: translateY(-10px); } 100% { transform: translateY(0px); }
    }
    .logo-container { display: flex; justify-content: center; margin-bottom: 10px; }
    .logo-img { width: 120px; border-radius: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.5); animation: float 3s ease-in-out infinite; }


    .stTextInput > div > div > input, .stSelectbox > div > div > div {
        background-color: #1e1e1e !important; color: #ededed !important; border: 1px solid #333 !important;
    }
    

    .info-box {
        background-color: rgba(211, 160, 235, 0.05);

        padding: 15px;
        border-radius: 5px;
        margin-bottom: 20px;
        color: #ededed;
    }


    .dep-tag {
        display: inline-block;
        background-color: #252525;

        padding: 4px 8px;
        border-radius: 4px;
        margin: 2px;
        font-size: 0.85rem;
        border: 1px solid #333;
    }

    #MainMenu {visibility: hidden;} footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)



def load_logo():

    if os.path.exists("assets/logo.png"):
        try:
            img = Image.open("assets/logo.png")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            data = base64.b64encode(buf.getvalue()).decode("utf-8")
            return f"data:image/png;base64,{data}"
        except: pass
    elif os.path.exists("logo.png"):
        try:
            img = Image.open("logo.png")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            data = base64.b64encode(buf.getvalue()).decode("utf-8")
            return f"data:image/png;base64,{data}"
        except: pass
    return None

def estimate_tokens(text):
    return len(text) // 4

def get_file_stats(files_dict):
    ext_counts = {}
    total_lines = 0
    total_size = 0
    for name, content in files_dict.items():
        ext = name.split('.')[-1] if '.' in name else 'folder/other'
        ext_counts[ext] = ext_counts.get(ext, 0) + 1
        lines = len(content.splitlines())
        total_lines += lines
        total_size += len(content)
    return ext_counts, total_lines, total_size

def calculate_github_quality_score(meta, commits, pr_stats):
    """
    Advanced Grading Logic based on GitHub Metrics.
    Returns (score, breakdown_dict)
    """
    score = 0
    breakdown = {}
    
    # 1. Popularity (30 pts)
    stars = meta.get('stars', 0)
    forks = meta.get('forks', 0)
    
    # Stars Logic (20 pts)
    if stars > 10000: star_pts = 20
    elif stars > 1000: star_pts = 18
    elif stars > 100: star_pts = 14
    elif stars > 10: star_pts = 8
    else: star_pts = 2
    
    # Fork Logic (10 pts) - Fork Ratio
    fork_ratio = forks / stars if stars > 0 else 0
    if fork_ratio > 0.2: fork_pts = 10
    elif fork_ratio > 0.1: fork_pts = 7
    else: fork_pts = 3
    
    popularity_score = star_pts + fork_pts
    score += popularity_score
    breakdown['Popularity'] = f"{popularity_score}/30 (Stars: {star_pts}, Fork Ratio: {fork_pts})"

    # 2. Activity (25 pts)
    # Recency
    last_commit_str = meta.get('pushed_at', None)
    activity_score = 0
    recency_pts = 0
    if last_commit_str:
        try:
            last_date = datetime.strptime(last_commit_str, "%Y-%m-%dT%H:%M:%SZ")
            days_diff = (datetime.utcnow() - last_date).days
            
            if days_diff < 30: recency_pts = 15
            elif days_diff < 180: recency_pts = 10
            elif days_diff < 365: recency_pts = 5
            else: recency_pts = 0
            
            activity_score += recency_pts
        except: pass
    
    # Commit Frequency (10 pts)
    freq_pts = 0
    if len(commits) > 50: freq_pts = 10
    elif len(commits) > 10: freq_pts = 7
    else: freq_pts = 3
    
    activity_score += freq_pts
    score += activity_score
    breakdown['Activity'] = f"{activity_score}/25 (Recency: {recency_pts}, Freq: {freq_pts})"

    # 3. Maintenance (20 pts)
    # Issues Heuristic
    open_issues = meta.get('open_issues', 0)
    issue_ratio = open_issues / stars if stars > 0 else 0
    
    if issue_ratio < 0.05: maint_pts = 20
    elif issue_ratio < 0.1: maint_pts = 15
    elif issue_ratio < 0.2: maint_pts = 10
    else: maint_pts = 5
    
    # Adjust based on PR Health
    pr_merge_rate = pr_stats.get('merge_rate', 0)
    if pr_merge_rate > 0.8: maint_pts += 5 # Bonus for healthy PRs
    elif pr_merge_rate < 0.2 and pr_stats.get('total_prs', 0) > 10: maint_pts -= 5 # Penalty for ignored PRs
    
    score += maint_pts
    breakdown['Maintenance'] = f"{maint_pts}/20 (Issue & PR Health)"

    # 4. Community (15 pts)
    watchers = meta.get('watchers', 0) 
    if watchers > 1000: comm_pts = 15
    elif watchers > 100: comm_pts = 10
    elif watchers > 10: comm_pts = 5
    else: comm_pts = 2
    
    score += comm_pts
    breakdown['Community'] = f"{comm_pts}/15 (Watchers)"

    # 5. Stability (10 pts)
    if meta.get('archived', False): stab_pts = 0
    else: stab_pts = 10
    
    score += stab_pts
    breakdown['Stability'] = f"{stab_pts}/10 (Active Status)"

    return score, breakdown

def get_grade_from_score(score):
    if score >= 95: return 'A++', "Exceptional quality, highly active, and massive community trust."
    elif score >= 90: return 'A+', "Excellent project with strong metrics and maintenance."
    elif score >= 80: return 'A', "Great project, reliable and well-maintained."
    elif score >= 70: return 'B+', "Good project, but might lack in activity or popularity."
    elif score >= 60: return 'B', "Fair quality, check specific metrics for details."
    elif score >= 50: return 'C+', "Average project, potential maintenance or activity issues."
    else: return 'C', "Low score, use with caution. May be inactive or unmaintained."

def analyze_static_quality(files_dict, total_lines):
    """Fallback for uploaded files (no GitHub API data)."""
    score = 50 
    reasons = []
    
    has_readme = any('readme' in f.lower() for f in files_dict.keys())
    if has_readme: 
        score += 20
        reasons.append("✅ README found (+20)")
    else: 
        reasons.append("❌ No README (-0)")
        
    total_comments = 0
    for content in files_dict.values():
        total_comments += len(re.findall(r'#.*|//.*|/\*.*?\*/', content, re.DOTALL))
    
    doc_ratio = total_comments / total_lines if total_lines > 0 else 0
    if doc_ratio > 0.2:
        score += 15
        reasons.append("✅ Great documentation (+15)")
    elif doc_ratio > 0.1:
        score += 10
        reasons.append("👍 Good documentation (+10)")
    else:
        reasons.append("⚠️ Low documentation (-0)")
        
    avg_lines = total_lines / len(files_dict) if files_dict else 0
    if avg_lines < 200:
        score += 15
        reasons.append("✅ Modular file structure (+15)")
    elif avg_lines > 1000:
        score -= 10
        reasons.append("❌ Bloated files (-10)")
        
    grade, desc = get_grade_from_score(score)
    return grade, desc, reasons

def build_full_tree(files_dict):
    lines = ["PROJECT ARCHITECTURE", "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"]
    sorted_files = sorted(files_dict.keys())
    has_main = any('main' in f.lower() or 'app' in f.lower() for f in sorted_files)
    
    lines.append("│")
    if has_main:
        lines.append("├── 🚀 Entry point detected")
    
    lines.append("│\n├── 📁 File Structure")
    
    for f in sorted_files:
        parts = f.split('/')
        depth = len(parts) - 1
        indent = "│   " * depth
        name = parts[-1]
        lines.append(f"{indent}├── 📄 {name} ({len(files_dict[f].splitlines())} lines)")
        
    return "\n".join(lines)

class StringIO:
    def __init__(self): self.buffer = io.StringIO()
    def write(self, text): self.buffer.write(text)
    def get_value(self): return self.buffer.getvalue()

def create_txt_content(files_dict):
    output = StringIO()
    output.write("RIKAI CODE EXPORT\n")
    output.write("==========================================\n\n")
    output.write(f"Note: The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}\n\n")
    output.write(build_full_tree(files_dict) + "\n\n")
    output.write("FILE CONTENTS\n")
    output.write("==========================================\n\n")
    for filename, content in files_dict.items():
        output.write(f"\n\n----- START: {filename} -----\n\n")
        output.write(content)
        output.write(f"\n\n----- END: {filename} -----\n")
    return output.get_value()

def export_html(files_dict):
    html = "<html><head><style>body{font-family:monospace; background:#171717; color:#ededed;} .file{margin-bottom:20px; border-bottom:1px solid #333;} h3{color:#d3a0eb;}</style></head><body>"
    html += f"<h1>RikaiCode Export</h1>"
    html += f"<p><strong>Note:</strong> The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}</p>"
    html += f"<pre>{build_full_tree(files_dict)}</pre><hr>"
    for k, v in files_dict.items():
        html += f"<div class='file'><h3>{k}</h3><pre>{v}</pre></div>"
    html += "</body></html>"
    return html


def export_docx(files_dict):
    doc = Document()
    doc.add_heading('RikaiCode Repository Export', 0)
    doc.add_paragraph(f"Note: The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}")
    doc.add_paragraph(build_full_tree(files_dict))
    for filename, content in files_dict.items():
        doc.add_heading(filename, level=1)
        try: doc.add_paragraph(content)
        except: doc.add_paragraph("[Content error]")
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def export_pdf(files_dict):
    class PDF(FPDF):
        def __init__(self):
            super().__init__()
            self.logo_path = LOGO_PATH
        
        def footer(self):
       
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            

            gen_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            self.cell(0, 10, f'Generated on {gen_date}', 0, 0, 'R')
            
            
            if os.path.exists(self.logo_path):

                self.image(self.logo_path, x=10, y=self.h - 12, h=5)

                self.set_xy(17, self.h - 12)
                self.cell(0, 5, 'RikaiCode', 0, 0, 'L')
            else:
                self.set_x(10)
                self.cell(0, 10, 'RikaiCode', 0, 0, 'L')

    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="RikaiCode Repository Export", ln=1, align='C')
    pdf.ln(5)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 5, f"Note: The following file types were excluded from detailed content processing to optimize performance: {', '.join(SKIP_EXTENSIONS)}")
    pdf.ln(5)
    
    for filename, content in files_dict.items():
        pdf.set_font("Arial", 'B', size=14)
        pdf.cell(200, 10, txt=filename, ln=1)
        pdf.set_font("Arial", size=10)
        try:
    
            if len(content) > 5000:
                safe_content = content[:5000] + "\n... [TRUNCATED FOR PDF PREVIEW]"
            else:
                safe_content = content
                
            safe_content = safe_content.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 5, safe_content)
            pdf.ln(5)
        except: pass
    

    pdf_bytes = pdf.output(dest='S')
    

    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode('latin-1')
        
    file_stream = io.BytesIO(pdf_bytes)
    file_stream.seek(0)
    return file_stream

def export_json(files_dict):
    data = {
        "metadata": { 
            "generated_at": datetime.now().isoformat(), 
            "file_count": len(files_dict), 
            "architecture": build_full_tree(files_dict),
            "excluded_types": list(SKIP_EXTENSIONS)
        },
        "files": [{"name": k, "content": v} for k, v in files_dict.items()]
    }
    return json.dumps(data, indent=2)



def scan_dependencies(files_dict):
    """Scans for dependencies in requirements.txt, package.json, or imports."""
    deps = set()
    
   
    if 'requirements.txt' in files_dict:
        for line in files_dict['requirements.txt'].splitlines():
            if line.strip() and not line.startswith('#'):
                pkg = line.strip().split('==')[0].split('>=')[0].split('<')[0]
                deps.add(f"🐍 {pkg}")
    
    if 'package.json' in files_dict:
        try:
            data = json.loads(files_dict['package.json'])
            for pkg in data.get('dependencies', {}).keys():
                deps.add(f"📦 {pkg}")
            for pkg in data.get('devDependencies', {}).keys():
                deps.add(f"🛠️ {pkg} (dev)")
        except: pass

    
    py_files = [f for f in files_dict if f.endswith('.py')][:5] 
    for f in py_files:
        matches = re.findall(r'^(?:import|from)\s+([a-zA-Z0-9_]+)', files_dict[f], re.MULTILINE)
        for m in matches:
            if m not in ['os', 'sys', 're', 'json', 'time', 'math']: # Filter stdlib
                deps.add(f"🐍 {m}")

    return sorted(list(deps))

def extract_code_structure(files_dict):
    """Extracts classes and functions for a summary view."""
    structure = {}
   
    target_files = [f for f in files_dict if f.endswith(('.py', '.js', '.ts'))] 
    
    for filename in target_files[:200]: # Limit to 200 files for performance
        content = files_dict[filename]
        

        if not content: continue
        
        classes = re.findall(r'class\s+([A-Za-z0-9_]+)', content)
        funcs = re.findall(r'def\s+([A-Za-z0-9_]+)', content)
        
        if classes or funcs:
            structure[filename] = {
                "classes": classes,
                "functions": funcs
            }
            
      
        del content
        
    return structure



def process_uploaded_files(uploaded_files):
    files_dict = {}
    for f in uploaded_files:
        try:
      
            ext = os.path.splitext(f.name)[1].lower()
            if ext in SKIP_EXTENSIONS:
                files_dict[f.name] = ""
            else:
                content = f.read().decode('utf-8')
                files_dict[f.name] = content
     
            del content
        except: pass
    gc.collect() 
    return files_dict

def process_zip_file(zip_file):
    files_dict = {}
    with zipfile.ZipFile(zip_file) as z:
        for filename in z.namelist():
            if filename.endswith('/'): continue
            
            ext = os.path.splitext(filename)[1].lower()
            clean_name = filename.split('/', 1)[-1] if '/' in filename else filename
            
            if ext in SKIP_EXTENSIONS:
                files_dict[clean_name] = "" 
                continue

            try:
                with z.open(filename) as f:
                    content = f.read().decode('utf-8')
                    files_dict[clean_name] = content
                    del content 
            except: pass
    gc.collect() 
    return files_dict

def process_github_url(url):
    files_dict = {}
    repo_meta = {}
    pr_stats = {}
    
    if url.endswith('.git'): url = url[:-4]
    
    try:
        parts = url.rstrip('/').split('/')
        owner, repo = parts[-2], parts[-1]
    except:
        st.error("Invalid GitHub URL format.")
        return {}, {}, {}

    api_base = f"https://api.github.com/repos/{owner}/{repo}"
    
    progress_bar = st.progress(0, text="Initializing...")
    
    try:
        progress_bar.progress(10, text="📡 Fetching repository metadata...")
        meta_r = requests.get(api_base)
        if meta_r.status_code == 200:
            data = meta_r.json()
            repo_meta['stars'] = data.get('stargazers_count', 0)
            repo_meta['forks'] = data.get('forks_count', 0)
            repo_meta['watchers'] = data.get('subscribers_count', data.get('watchers_count', 0)) 
            repo_meta['open_issues'] = data.get('open_issues_count', 0)
            repo_meta['language'] = data.get('language', 'Unknown')
            repo_meta['archived'] = data.get('archived', False)
            repo_meta['pushed_at'] = data.get('pushed_at', None)
            repo_meta['created_at'] = data.get('created_at', None)
            

            if repo_meta['created_at']:
                created = datetime.strptime(repo_meta['created_at'], "%Y-%m-%dT%H:%M:%SZ")
                age_days = (datetime.utcnow() - created).days
                repo_meta['age_years'] = round(age_days / 365, 1)
        else:
            progress_bar.empty()
            st.error(f"GitHub API Error: {meta_r.status_code}")
            return {}, {}, {}

        
        progress_bar.progress(15, text="📡 Analyzing Pull Requests...")
        try:
            r_open = requests.get(f"https://api.github.com/search/issues?q=repo:{owner}/{repo}+type:pr+state:open")
            open_prs = r_open.json().get('total_count', 0) if r_open.status_code == 200 else 0
            
            r_merged = requests.get(f"https://api.github.com/search/issues?q=repo:{owner}/{repo}+type:pr+is:merged")
            merged_prs = r_merged.json().get('total_count', 0) if r_merged.status_code == 200 else 0
            
            r_closed = requests.get(f"https://api.github.com/search/issues?q=repo:{owner}/{repo}+type:pr+state:closed")
            closed_total = r_closed.json().get('total_count', 0) if r_closed.status_code == 200 else 0
            
            pr_stats['open'] = open_prs
            pr_stats['merged'] = merged_prs
            pr_stats['closed_rejected'] = max(0, closed_total - merged_prs)
            pr_stats['total_prs'] = open_prs + closed_total
            
            if closed_total > 0:
                pr_stats['merge_rate'] = merged_prs / closed_total
            else:
                pr_stats['merge_rate'] = 0
                
        except Exception as e:
            st.warning(f"Could not fetch detailed PR stats: {e}")

    
        progress_bar.progress(20, text="📡 Fetching commit history...")
        commits_r = requests.get(f"{api_base}/commits?per_page=100")
        commit_dates = []
        if commits_r.status_code == 200:
            for c in commits_r.json():
                try:
                    date_str = c['commit']['author']['date']
                    d = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%SZ").date()
                    commit_dates.append(d)
                except: pass
            repo_meta['commit_dates'] = commit_dates


        archive_url_main = f"{api_base}/zipball/main"
        archive_url_master = f"{api_base}/zipball/master"
        
        progress_bar.progress(30, text="🔍 Checking Main branch...")
        r = requests.get(archive_url_main, stream=True)
        
        if r.status_code == 404:
            progress_bar.progress(40, text="🔍 Main not found. Checking Master...")
            r = requests.get(archive_url_master, stream=True)
        
        if r.status_code != 200:
            progress_bar.empty()
            st.error("Could not fetch repository content.")
            return {}, {}, {}

        total_size = int(r.headers.get('content-length', 0))
        chunk_data = []
        downloaded = 0
        
        progress_bar.progress(50, text=f"⬇️ Downloading ({int(total_size/1024)} KB)...")
        
        for chunk in r.iter_content(chunk_size=8192):
            if chunk:
                chunk_data.append(chunk)
                downloaded += len(chunk)
                if total_size > 0:
                    pct = 50 + int((downloaded / total_size) * 30)
                    progress_bar.progress(pct, text=f"⬇️ Downloading... {int(downloaded/1024)}/{int(total_size/1024)} KB")
        
        progress_bar.progress(80, text="📦 Decompressing...")
        zip_data = io.BytesIO(b''.join(chunk_data))
        
        with zipfile.ZipFile(zip_data) as z:
            file_list = z.namelist()
            total_files = len(file_list)
            processed_count = 0
            
            for filename in file_list:
                if filename.endswith('/'): continue
                
              
                ext = os.path.splitext(filename)[1].lower()
                clean_name = filename.split('/', 1)[-1] if '/' in filename else filename

                if ext in SKIP_EXTENSIONS:
                    files_dict[clean_name] = "" 
                else:
                    try:
                        with z.open(filename) as f:
                            content = f.read().decode('utf-8')
                            files_dict[clean_name] = content
                            del content 
                    except: pass
                
                processed_count += 1
                if total_files > 0:
                    pct = 80 + int((processed_count / total_files) * 20)
                    progress_bar.progress(pct, text=f"⚙️ Processing {processed_count}/{total_files}")
        
        progress_bar.progress(100, text="✔️ Done!")
        time.sleep(0.5)
        progress_bar.empty()
        
        
        del chunk_data
        del zip_data
        gc.collect()
        
    except Exception as e:
        progress_bar.empty()
        st.error(f"Error: {str(e)}")
        
    return files_dict, repo_meta, pr_stats


if 'files_data' not in st.session_state: st.session_state.files_data = {}
if 'repo_meta' not in st.session_state: st.session_state.repo_meta = {}
if 'pr_stats' not in st.session_state: st.session_state.pr_stats = {}
if 'search_term' not in st.session_state: st.session_state.search_term = ""


logo_data = load_logo()
if logo_data:
    st.markdown(f'<div class="logo-container"><img src="{logo_data}" class="logo-img"></div>', unsafe_allow_html=True)

st.markdown("""
<h1 style='text-align: center; margin-bottom: 0;'>
    Rikai<span style='color: #d3a0eb;'>Code</span>
</h1>
<p style='text-align: center; font-size: 1.2rem; color: #a0a0a0; margin-bottom: 30px;'>
    Advanced Repository Flattener, Context Generator & Intelligent Grader
</p>
""", unsafe_allow_html=True)







st.markdown("### 🛠️ Input Source")
input_method = st.selectbox(
    "Select source type", 
    ["🌐 GitHub Repository URL", "📁 Upload Files (Multi-Select)", "📦 Upload ZIP Folder"],
    label_visibility="collapsed"
)


if 'prev_input_method' not in st.session_state:
    st.session_state.prev_input_method = input_method

if st.session_state.prev_input_method != input_method:
    st.session_state.files_data = {}
    st.session_state.repo_meta = {}
    st.session_state.pr_stats = {}
    st.session_state.prev_input_method = input_method
    gc.collect()

if input_method == "📁 Upload Files (Multi-Select)":
    uploaded_files = st.file_uploader("Drag and drop files", accept_multiple_files=True, key="file_uploader_widget")
    if uploaded_files:
        st.session_state.files_data = process_uploaded_files(uploaded_files)
        st.session_state.repo_meta = {} 
        st.session_state.pr_stats = {}

elif input_method == "📦 Upload ZIP Folder":
    zip_file = st.file_uploader("Drag and drop ZIP", type=['zip'], key="zip_uploader_widget")
    if zip_file:
        st.session_state.files_data = process_zip_file(zip_file)
        st.session_state.repo_meta = {}
        st.session_state.pr_stats = {}

elif input_method == "🌐 GitHub Repository URL":
    url = st.text_input("Enter Public GitHub URL", placeholder="https://github.com/user/repo", key="github_url_input")
    if st.button("Fetch Repository", key="fetch_github_btn"):
        if url:
          
            st.session_state.files_data = {}
            st.session_state.repo_meta = {}
            st.session_state.pr_stats = {}
            gc.collect()
            
            with st.spinner("Loading large repository… Please wait while the architecture loads. ⏳"):
                files, meta, pr_stats = process_github_url(url)
                st.session_state.files_data = files
                st.session_state.repo_meta = meta
                st.session_state.pr_stats = pr_stats



st.markdown(f"""
<div class="info-box">
    <p><strong>Note:</strong> Large repositories may take time to process. Please wait for the architecture to load.</p>
    <p><strong>Local Files:</strong> If uploading files/ZIPs manually, GitHub stats and grading will be estimated based on static analysis only.</p>
    <p><strong>Optimization:</strong> To reduce processing time, the following file types are excluded from content analysis but listed in architecture: <strong>{', '.join(SKIP_EXTENSIONS)}</strong>.</p>
</div>
""", unsafe_allow_html=True)


if st.session_state.files_data:
    files_data = st.session_state.files_data
    repo_meta = st.session_state.repo_meta
    pr_stats = st.session_state.pr_stats
    
    st.markdown("---")
    

    if repo_meta:
        st.markdown("### 🌐 Repository Insights")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("⭐ Stars", repo_meta.get('stars', 0), help="Total number of stars the repository has received.")
        c2.metric("🍴 Forks", repo_meta.get('forks', 0), help="Total number of forks created from this repository.")
        c3.metric("👀 Watchers", repo_meta.get('watchers', 0), help="Number of users watching the repository for updates.")
        c4.metric("💻 Language", repo_meta.get('language', 'N/A'), help="The primary programming language used in the repository.")
        
    
        st.markdown("#### ~ Derived Metrics")
        dm1, dm2, dm3, dm4 = st.columns(4)
        
        
        fork_ratio = round(repo_meta.get('forks', 0) / repo_meta.get('stars', 1), 2) if repo_meta.get('stars', 0) > 0 else 0
        dm1.metric("🍴 Fork Ratio", f"{fork_ratio}", help="Forks per Star. High ratio implies high utility/reuse.")
        
  
        age = repo_meta.get('age_years', 'N/A')
        dm2.metric("📅 Repo Age", f"{age} years", help="Time since repository creation.")
        
     
        open_i = repo_meta.get('open_issues', 0)
        dm3.metric("❓ Open Issues", f"{open_i}", help="Count of open issues. Compare with activity to gauge maintenance load.")
        

        merge_rate = pr_stats.get('merge_rate', 0)
        dm4.metric("🔄 PR Merge Rate", f"{merge_rate:.0%}", help="Percentage of closed PRs that were merged. High rate = healthy contribution flow.")


        st.markdown("#### ~ Pull Request Deep Dive")
        pc1, pc2, pc3, pc4 = st.columns(4)
        pc1.metric("Open PRs", pr_stats.get('open', 'N/A'), help="Pull Requests currently open and awaiting review or merge.")
        pc2.metric("Merged PRs", pr_stats.get('merged', 'N/A'), help="Pull Requests that have been successfully merged into the codebase.")
        pc3.metric("Closed (Rejected)", pr_stats.get('closed_rejected', 'N/A'), help="Pull Requests that were closed without being merged.")
        pc4.metric("📖 Total PRs", pr_stats.get('total_prs', 'N/A'), help="Total number of Pull Requests created in the repository.")
        
        st.markdown("---")
        if 'commit_dates' in repo_meta and repo_meta['commit_dates']:
            st.markdown("#### ~ Recent Commit Activity")
            df_commits = pd.DataFrame(repo_meta['commit_dates'], columns=['date'])
            df_commits['count'] = 1
            df_commits = df_commits.groupby('date').sum().reset_index()
            
            fig_commits = px.bar(df_commits, x='date', y='count', 
                                 labels={'date': 'Date', 'count': 'Commits'},
                                 color_discrete_sequence=['#d3a0eb'])
            fig_commits.update_layout(
                paper_bgcolor='#171717', plot_bgcolor='#171717', 
                font_color='#ededed', xaxis=dict(gridcolor='#333'), yaxis=dict(gridcolor='#333')
            )
            st.plotly_chart(fig_commits, width="stretch")
        st.markdown("---")


    stats, total_lines, total_size = get_file_stats(files_data)
    content_text = create_txt_content(files_data)
    token_est = estimate_tokens(content_text)
    

    if repo_meta:
        score, breakdown = calculate_github_quality_score(repo_meta, repo_meta.get('commit_dates', []), pr_stats)
        grade, grade_desc = get_grade_from_score(score)
        
        col_grade, col_stats = st.columns([1, 3])
        
        with col_grade:
            st.markdown(f"""
            <div class="grade-box">
                <div style="font-size: 1.2rem; color: #ededed;">Repository Rank</div>
                <div class="grade-letter">{grade}</div>
                <div style="color: #a0a0a0;">{grade_desc}</div>
                <div class="grade-reason">
                    <strong>Score Breakdown:</strong><br>
                    • Popularity: {breakdown.get('Popularity', 'N/A')}<br>
                    • Activity: {breakdown.get('Activity', 'N/A')}<br>
                    • Maintenance: {breakdown.get('Maintenance', 'N/A')}<br>
                    • Community: {breakdown.get('Community', 'N/A')}<br>
                    • Stability: {breakdown.get('Stability', 'N/A')}
                </div>
            </div>
            """, unsafe_allow_html=True)
    else:
       
        grade, grade_desc, reasons = analyze_static_quality(files_data, total_lines)
        col_grade, col_stats = st.columns([1, 3])
        with col_grade:
            reason_text = "<br>".join(reasons)
            st.markdown(f"""
            <div class="grade-box">
                <div style="font-size: 1.2rem; color: #ededed;">Static Grade</div>
                <div class="grade-letter">{grade}</div>
                <div style="color: #a0a0a0;">{grade_desc}</div>
                <div class="grade-reason">
                    <strong>Static Analysis:</strong><br>
                    {reason_text}
                </div>
            </div>
            """, unsafe_allow_html=True)

    with col_stats:
        st.markdown("### ⌨️ Code Statistics")
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Total Files", len(files_data))
        m2.metric("Total Lines", f"{total_lines:,}")
        m3.metric("Characters", f"{total_size:,}")
        m4.metric("Est. Tokens", f"{token_est:,}")
        m5.metric("Est. Size", f"{total_size/1024:.1f} KB")

    
    is_huge_repo = total_lines > MAX_LINES_INTERACTIVE_PREVIEW
    if is_huge_repo:
        st.warning(f"⚠️ **Large Repository Detected ({total_lines:,} lines).** Interactive preview disabled to prevent crash.")


    st.markdown("---")
    st.markdown("### ~ Code Frame")
    
    ac1, ac2 = st.columns(2)
    
    with ac1:
        st.markdown("#### Detected Dependencies")
        with st.expander("View Dependencies", expanded=False):
            deps = scan_dependencies(files_data)
            if deps:
                deps_html = " ".join([f"<span class='dep-tag'>{d}</span>" for d in deps])
                st.markdown(deps_html, unsafe_allow_html=True)
            else:
                st.info("No external dependencies found in common files.")
    
    with ac2:
        st.markdown("####  Code Structure (Classes/Funcs)")
        with st.expander("View Structure", expanded=False):
            structure = extract_code_structure(files_data)
            if structure:
                for fname, items in structure.items():
                    st.markdown(f"**{fname}**")
                    if items['classes']:
                        classes_str = ", ".join(items['classes'])
                        st.markdown(f"▫️ Classes: `{classes_str}`")
                    if items['functions']:
                        funcs_str = ", ".join(items['functions'])
                        st.markdown(f"▫️ Functions: `{funcs_str}`")
                    st.markdown("---")
            else:
                st.info("No structure detected or files too large.")


    st.markdown("---")
    st.markdown("### 📉 File Distribution Analysis")
    

    color_seq = px.colors.qualitative.Vivid
    

    df = pd.DataFrame(list(stats.items()), columns=['Extension', 'Count'])
    
    col_treemap, col_pie = st.columns(2)
    
    with col_treemap:
        st.markdown("#### Treemap View")
        fig_tree = px.treemap(df, path=['Extension'], values='Count', 
                         color='Count', color_continuous_scale='Viridis')
        fig_tree.update_layout(
            paper_bgcolor='#171717', font_color='#ededed',
            plot_bgcolor='#171717', margin=dict(t=0, l=0, r=0, b=0)
        )
        st.plotly_chart(fig_tree, width="stretch")

    with col_pie:
        st.markdown("#### Distribution View")
        fig_pie = px.pie(df, values='Count', names='Extension', hole=0.4, color_discrete_sequence=color_seq)
        fig_pie.update_layout(
            paper_bgcolor='#171717', font_color='#ededed',
            plot_bgcolor='#171717', legend=dict(bgcolor='#171717', font=dict(color='#ededed'))
        )
        st.plotly_chart(fig_pie, width="stretch")


    st.markdown("---")
    st.markdown("### ~ Project Architecture")
    arch_str = build_full_tree(files_data)
    
    if len(arch_str.splitlines()) > 30:
        with st.expander("View Full Architecture Diagram", expanded=False):
            st.markdown(f'<div class="arch-box">{arch_str}</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="arch-box">{arch_str}</div>', unsafe_allow_html=True)

  
    if not is_huge_repo:
        st.markdown("---")
        st.markdown("### 📂 Code Preview")
        
        def update_search(): st.session_state.search_term = st.session_state.search_input
        search_val = st.text_input(
    "",
    value=st.session_state.search_term,
    key="search_input",
    on_change=update_search,
    placeholder="Enter a file name to search...",
    label_visibility="collapsed"
)
        
        for filename, content in files_data.items():
            if st.session_state.search_term.lower() not in filename.lower():
                continue
                
            ext = filename.split('.')[-1]
            icon = "🐍" if ext == 'py' else "📜"
            lines_count = len(content.splitlines())
            
           
            single_file_data = content.encode('utf-8')
            
            with st.expander(f"{icon} **{filename}** ({lines_count} lines)"):
                c_prev, c_dl = st.columns([4, 1])
                with c_prev:
                    st.code(content, language='python' if ext=='py' else 'javascript')
                with c_dl:
                    st.download_button(
                        label="Download",
                        data=single_file_data,
                        file_name=filename,
                        mime="text/plain",
                        key=f"dl_single_{filename}"
                    )


    st.markdown("---")
    st.markdown("### 🔗 Export Options")
    
  
    e1, e2, e3, e4, e5, e6, e7 = st.columns(7)
    
    e1.download_button("Export Text (.txt)", content_text, "RikaiCode.txt", "text/plain")
    e5.download_button("Export Markdown (.md)", f"```\n{arch_str}\n```\n\n" + content_text, "RikaiCode.md", "text/markdown")
    e2.download_button("Export JSON (.json)", export_json(files_data), "RikaiCode.json", "application/json")
    e3.download_button("Export HTML (.html)", export_html(files_data), "RikaiCode.html", "text/html")
    
    latex_content = f"\\documentclass{{article}}\\begin{{document}}\\begin{{verbatim}}{arch_str}\\end{{verbatim}}\\end{{document}}"
    e4.download_button("Export LaTeX (.tex)", latex_content, "RikaiCode.tex", "application/x-tex")
    
    if e6.button("Generate DOCX"):
        with st.spinner("Generating..."):
            docx_file = export_docx(files_data)
            e6.download_button("Download DOCX", docx_file, "RikaiCode.docx")

    if e7.button("Generate PDF"):
        with st.spinner("Generating..."):
            pdf_file = export_pdf(files_data)
            e7.download_button("Download PDF", pdf_file, "RikaiCode.pdf")


gc.collect()

st.markdown("""
<div class="footer-custom">
    <div>Made with 🤍  by <strong>Aurumz</strong></div>
    <div><a href="https://github.com/aurumz-rgb/RikaiCode" target="_blank" class="footer-link">© 2026 aurumz-rgb — AGPL 3.0 License</a></div>
</div>
""", unsafe_allow_html=True)